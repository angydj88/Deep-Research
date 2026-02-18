"""Generador de reportes en múltiples formatos."""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config import Config


class ReportBuilder:
    """Construye reportes en Markdown, DOCX y TXT."""

    def __init__(self, objetivo: str):
        self.objetivo = objetivo
        self.secciones: list[dict] = []
        self.resumen = ""
        self.timestamp = datetime.now()

        # Crear directorio de reportes
        os.makedirs(Config.REPORTS_DIR, exist_ok=True)

    def agregar_seccion(self, dimension: dict, contenido: str,
                        fuentes: list[str], exito: bool):
        """Agrega una sección completada."""
        self.secciones.append({
            "dimension": dimension,
            "contenido": contenido,
            "fuentes": fuentes,
            "exito": exito,
        })

    def set_resumen(self, resumen: str):
        self.resumen = resumen

    # ── Estadísticas ──
    @property
    def stats(self) -> dict:
        exitosas = sum(1 for s in self.secciones if s["exito"])
        total_chars = sum(len(s["contenido"]) for s in self.secciones)
        return {
            "exitosas": exitosas,
            "total": len(self.secciones),
            "caracteres": total_chars,
        }

    # ── Exportar Markdown ──
    def exportar_markdown(self) -> str:
        """Genera el informe completo en Markdown."""

        lineas = [
            f"# 🔬 MEGA INFORME DE INVESTIGACIÓN PROFUNDA",
            f"",
            f"**Objetivo:** {self.objetivo}",
            f"**Fecha:** {self.timestamp.strftime('%d/%m/%Y %H:%M')}",
            f"**Modelo:** {Config.MODEL}",
            f"**Versión:** Deep Research v{Config.VERSION}",
            f"",
            f"---",
            f"",
            f"## 📑 Tabla de Contenidos",
            f"",
        ]

        for s in self.secciones:
            d = s["dimension"]
            lineas.append(
                f"{d['num']}. {d['emoji']} {d['nombre']}"
            )

        lineas.extend(["", "---", ""])

        for s in self.secciones:
            d = s["dimension"]
            lineas.extend([
                f"## {d['emoji']} Sección {d['num']}: {d['nombre']}",
                f"",
                s["contenido"],
                f"",
            ])

            if s["fuentes"]:
                lineas.append("### 📚 Fuentes")
                for f in s["fuentes"]:
                    lineas.append(f"- {f}")
                lineas.append("")

            lineas.extend(["---", ""])

        if self.resumen:
            lineas.extend([
                "## 📋 RESUMEN EJECUTIVO INTEGRADO",
                "",
                self.resumen,
                "",
                "---",
                "",
            ])

        stats = self.stats
        lineas.extend([
            "## 📊 Metadata",
            f"- Secciones exitosas: {stats['exitosas']}/{stats['total']}",
            f"- Caracteres totales: {stats['caracteres']:,}",
            f"- Modelo: {Config.MODEL}",
            f"- Generado: {self.timestamp.strftime('%d/%m/%Y %H:%M')}",
        ])

        return "\n".join(lineas)

    # ── Exportar TXT (para alimentar IA) ──
    def exportar_texto_plano(self) -> str:
        """Texto limpio optimizado para LLMs."""
        partes = [
            "═" * 50,
            "MEGA INFORME DE INVESTIGACIÓN PROFUNDA",
            "═" * 50,
            "",
            f"OBJETIVO: {self.objetivo}",
            f"FECHA: {self.timestamp.strftime('%d/%m/%Y %H:%M')}",
            "",
        ]

        for s in self.secciones:
            d = s["dimension"]
            partes.extend([
                "─" * 50,
                f"SECCIÓN {d['num']}: {d['nombre']}",
                "─" * 50,
                "",
                s["contenido"],
                "",
            ])

        if self.resumen:
            partes.extend([
                "─" * 50,
                "RESUMEN EJECUTIVO",
                "─" * 50,
                "",
                self.resumen,
            ])

        partes.extend(["", "═" * 50, "FIN DEL INFORME", "═" * 50])

        return "\n".join(partes)

    # ── Exportar DOCX ──
    def exportar_docx(self) -> str:
        """Genera archivo Word y retorna la ruta."""
        doc = Document()

        # Portada
        titulo = doc.add_paragraph()
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = titulo.add_run("MEGA INFORME DE INVESTIGACIÓN PROFUNDA")
        run.bold = True
        run.font.size = Pt(20)

        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.add_run(self.objetivo).font.size = Pt(14)

        doc.add_paragraph(
            f"Fecha: {self.timestamp.strftime('%d/%m/%Y')} | "
            f"Modelo: {Config.MODEL}"
        ).alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

        # Secciones
        for s in self.secciones:
            d = s["dimension"]
            doc.add_heading(
                f"{d['emoji']} {d['nombre']}", level=1
            )

            for parrafo in s["contenido"].split("\n"):
                parrafo = parrafo.strip()
                if not parrafo:
                    continue
                if parrafo.startswith("### "):
                    doc.add_heading(parrafo[4:], level=3)
                elif parrafo.startswith("## "):
                    doc.add_heading(parrafo[3:], level=2)
                elif parrafo.startswith("# "):
                    doc.add_heading(parrafo[2:], level=2)
                elif parrafo.startswith(("- ", "* ")):
                    doc.add_paragraph(parrafo[2:], style="List Bullet")
                else:
                    doc.add_paragraph(parrafo)

            doc.add_page_break()

        # Resumen
        if self.resumen:
            doc.add_heading("📋 RESUMEN EJECUTIVO", level=1)
            for p in self.resumen.split("\n"):
                if p.strip():
                    doc.add_paragraph(p)

        # Guardar
        nombre = self._nombre_archivo("docx")
        ruta = os.path.join(Config.REPORTS_DIR, nombre)
        doc.save(ruta)
        return ruta

    # ── Guardar archivos ──
    def guardar_todo(self) -> dict:
        """Guarda en todos los formatos y retorna rutas."""
        rutas = {}

        # Markdown
        md = self.exportar_markdown()
        ruta_md = os.path.join(
            Config.REPORTS_DIR, self._nombre_archivo("md")
        )
        with open(ruta_md, "w", encoding="utf-8") as f:
            f.write(md)
        rutas["markdown"] = ruta_md

        # Texto plano
        txt = self.exportar_texto_plano()
        ruta_txt = os.path.join(
            Config.REPORTS_DIR, self._nombre_archivo("txt")
        )
        with open(ruta_txt, "w", encoding="utf-8") as f:
            f.write(txt)
        rutas["texto"] = ruta_txt

        # DOCX
        rutas["docx"] = self.exportar_docx()

        return rutas

    def _nombre_archivo(self, extension: str) -> str:
        slug = self.objetivo[:40].replace(" ", "_")
        slug = "".join(
            c for c in slug if c.isalnum() or c == "_"
        )
        fecha = self.timestamp.strftime("%Y%m%d_%H%M")
        return f"informe_{slug}_{fecha}.{extension}"
